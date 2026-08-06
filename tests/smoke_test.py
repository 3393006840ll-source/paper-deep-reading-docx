from __future__ import annotations

import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pypdf import PdfWriter

ROOT = Path(__file__).resolve().parents[1]

REQUIRED_SECTIONS = [
    "五分钟快速筛选",
    "论文基本信息",
    "研究逻辑",
    "研究对象与工况",
    "研究方法",
    "核心公式与模型解释",
    "图表专项解读",
    "主要结果",
    "创新点",
    "局限性",
    "批判性评价",
    "研究方向/实际问题",
    "基础知识",
    "文献卡片",
    "最后浓缩",
    "三分钟导师汇报稿",
    "文献综述",
    "幻觉与证据核查",
]


def configure_table(table: object) -> None:
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        tr_pr.append(header)
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="paper_deep_reading_smoke_") as temp_dir:
        temp = Path(temp_dir)
        pdf_path = temp / "paper.pdf"
        evidence_path = temp / "evidence.txt"
        writer = PdfWriter()
        writer.add_blank_page(width=72, height=72)
        with pdf_path.open("wb") as handle:
            writer.write(handle)
        subprocess.run(
            [
                sys.executable,
                str(ROOT / "scripts" / "extract_pdf_evidence.py"),
                str(pdf_path),
                str(evidence_path),
            ],
            cwd=ROOT,
            check=True,
        )
        evidence = evidence_path.read_text(encoding="utf-8")
        assert "PDF_PAGES: 1" in evidence
        assert "===== PDF PAGE 1 =====" in evidence

        document = Document()
        document.core_properties.author = ""
        for section in REQUIRED_SECTIONS:
            document.add_heading(section, level=1)
            document.add_paragraph("论文未说明")
        table = document.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "证据位置"
        configure_table(table)
        docx_path = temp / "output.docx"
        document.save(docx_path)
        audit = subprocess.run(
            [
                sys.executable,
                str(ROOT / "scripts" / "audit_docx.py"),
                str(docx_path),
            ],
            cwd=ROOT,
            capture_output=True,
            text=True,
        )
        if audit.returncode != 0:
            raise AssertionError(audit.stdout + audit.stderr)

    print("PASSED: temporary PDF extraction and DOCX audit smoke test.")


if __name__ == "__main__":
    main()
