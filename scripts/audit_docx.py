#!/usr/bin/env python3
"""Audit structural requirements of a 研0 literature-reading DOCX."""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


REQUIRED_TEXT = [
    "五分钟快速筛选",
    "论文基本信息",
    "研究逻辑",
    "研究对象与工况",
    "研究方法",
    "核心公式与模型解释",
    "图表专项解读",
    "主要结果",
    "创新点",
    "局限性",
    "批判性评价",
    "培养方向",
    "基础知识",
    "文献卡片",
    "最后浓缩",
    "三分钟导师汇报稿",
    "文献综述",
    "幻觉与证据核查",
]


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument(
        "--allow-missing-sections",
        action="store_true",
        help="Report missing standard sections as warnings rather than errors.",
    )
    args = parser.parse_args()

    if not args.docx.is_file():
        parser.error(f"DOCX not found: {args.docx}")

    document = Document(args.docx)
    full_text = "\n".join(p.text for p in document.paragraphs)
    errors: list[str] = []
    warnings: list[str] = []

    if document.core_properties.author:
        errors.append("Document author metadata is not blank.")

    missing = [item for item in REQUIRED_TEXT if item not in full_text]
    if missing:
        target = warnings if args.allow_missing_sections else errors
        target.append("Missing standard sections: " + ", ".join(missing))

    if not document.tables:
        errors.append("No tables found.")

    for table_index, table in enumerate(document.tables, start=1):
        if not table._tbl.tblPr.xpath('./w:tblLayout[@w:type="fixed"]'):
            errors.append(f"Table {table_index} does not use fixed layout.")
        if table.rows and not table.rows[0]._tr.xpath("./w:trPr/w:tblHeader"):
            errors.append(f"Table {table_index} does not repeat its header row.")
        for row_index, row in enumerate(table.rows, start=1):
            if not row._tr.xpath("./w:trPr/w:cantSplit"):
                errors.append(
                    f"Table {table_index}, row {row_index} may split across pages."
                )

    if "论文未说明" not in full_text:
        warnings.append(
            "The document never uses “论文未说明”; confirm the paper truly reports all "
            "reproducibility and uncertainty details."
        )

    print(f"Paragraphs: {len(document.paragraphs)}")
    print(f"Tables: {len(document.tables)}")
    for warning in warnings:
        print(f"WARNING: {warning}")
    for error in errors:
        print(f"ERROR: {error}")

    if errors:
        print(f"FAILED: {len(errors)} structural issue(s).")
        return 1
    print("PASSED: structural audit complete; visual page inspection is still required.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
